import requests
import os
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

cur_path = os.path.dirname(os.path.abspath(__file__))

def html_formater(file_name, src='download', dst='story'):
    file_id = str(int(file_name.split("-")[1])-900+1)
    html_text = None
    with open(os.path.join(cur_path, src, file_name), 'r', encoding="utf-8") as fd:
        html_text = fd.read()
    if html_text:
        soup = BeautifulSoup(html_text, 'html.parser')
        story_div = soup.find("div", attrs={"class":"box_con"})
        page_title = story_div.find("div",attrs={"class":"bookname"}).h1.text
        page_content = story_div.find("div", id="content").strings
        with open(os.path.join(cur_path, dst, file_id + "-" + page_title+".txt"),"w",encoding='utf-8') as fd:
            fd.writelines(map(lambda x: x.strip() + "\n", page_content))

def html_formater_doc():
    downloaded_dir_name = "story"
    files = os.listdir(os.path.join(cur_path, downloaded_dir_name))
    files.sort(key=lambda x:int(x.split("-")[0]))
    doc = Document()
    doc.styles['Normal'].font.name='宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    for file in files:
        head = file.split("-")[1].split(".")[0]
        doc.add_heading(head,level=1)
        content = ""
        with open(os.path.join(cur_path, downloaded_dir_name, file),'r', encoding="utf-8") as fd:
            content = fd.read()
        p = doc.add_paragraph()
        p.add_run(content).font.size = Pt(18)
    doc.save("魔女之旅.docx")

if __name__ == "__main__":
    downloaded_dir_name = "download"
    files = os.listdir(os.path.join(cur_path, downloaded_dir_name))
    for file in files:
        try:
            html_formater(file, downloaded_dir_name)
        except:
            print(file + " format failed")
    html_formater_doc()